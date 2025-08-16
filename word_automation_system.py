import os
import json
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
from datetime import datetime
import win32com.client as win32
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
import logging
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class DocumentRequest:
    """Structure for document generation requests"""
    doc_type: str  # 'article', 'report', 'memo', 'presentation'
    topic: str
    length: str  # 'short', 'medium', 'long'
    tone: str  # 'formal', 'casual', 'technical'
    audience: str
    data_sources: Optional[List[str]] = None
    include_charts: bool = False
    template: Optional[str] = None

class AIContentGenerator:
    """Handles AI-powered content generation"""
    
    def __init__(self, api_key: str = None):
        self.api_key = api_key
        if api_key:
            openai.api_key = api_key
    
    def generate_content(self, request: DocumentRequest) -> Dict[str, Any]:
        """Generate content based on document request"""
        
        # Define content structure based on document type
        content_structure = self._get_content_structure(request.doc_type)
        
        # Generate content for each section
        content = {}
        for section in content_structure:
            prompt = self._create_prompt(section, request)
            content[section] = self._generate_text(prompt, request.length)
        
        return content
    
    def _get_content_structure(self, doc_type: str) -> List[str]:
        """Define document structure based on type"""
        structures = {
            'article': ['title', 'introduction', 'main_content', 'conclusion'],
            'report': ['executive_summary', 'introduction', 'methodology', 'findings', 'recommendations', 'conclusion'],
            'memo': ['header', 'purpose', 'summary', 'details', 'action_items'],
            'presentation': ['title_slide', 'agenda', 'main_points', 'conclusion', 'next_steps']
        }
        return structures.get(doc_type, ['title', 'content', 'conclusion'])
    
    def _create_prompt(self, section: str, request: DocumentRequest) -> str:
        """Create AI prompt for specific section"""
        base_prompt = f"""
        Write a {section} for a {request.doc_type} about {request.topic}.
        Tone: {request.tone}
        Audience: {request.audience}
        Length: {request.length}
        """
        return base_prompt
    
    def _generate_text(self, prompt: str, length: str) -> str:
        """Generate text using AI model (placeholder for actual AI integration)"""
        # This would integrate with OpenAI, Claude, or other AI models
        # For now, returning structured placeholder content
        length_words = {'short': 100, 'medium': 300, 'long': 500}
        target_length = length_words.get(length, 300)
        
        return f"[Generated content for: {prompt[:50]}...] (Target length: {target_length} words)"

class DataAnalyzer:
    """Handles data analysis and visualization"""
    
    def __init__(self):
        self.supported_formats = ['.csv', '.xlsx', '.json', '.xml']
    
    def analyze_data(self, data_path: str, analysis_type: str = 'summary') -> Dict[str, Any]:
        """Perform data analysis"""
        try:
            # Load data based on file format
            data = self._load_data(data_path)
            
            # Perform analysis
            analysis_results = {
                'summary': self._generate_summary(data),
                'insights': self._generate_insights(data),
                'recommendations': self._generate_recommendations(data)
            }
            
            return analysis_results
        
        except Exception as e:
            logger.error(f"Data analysis failed: {str(e)}")
            return {'error': str(e)}
    
    def _load_data(self, data_path: str) -> pd.DataFrame:
        """Load data from various formats"""
        file_ext = Path(data_path).suffix.lower()
        
        if file_ext == '.csv':
            return pd.read_csv(data_path)
        elif file_ext in ['.xlsx', '.xls']:
            return pd.read_excel(data_path)
        elif file_ext == '.json':
            return pd.read_json(data_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _generate_summary(self, data: pd.DataFrame) -> Dict[str, Any]:
        """Generate data summary"""
        return {
            'rows': len(data),
            'columns': len(data.columns),
            'numeric_columns': len(data.select_dtypes(include='number').columns),
            'missing_values': data.isnull().sum().to_dict(),
            'basic_stats': data.describe().to_dict()
        }
    
    def _generate_insights(self, data: pd.DataFrame) -> List[str]:
        """Generate data insights"""
        insights = []
        
        # Basic insights
        insights.append(f"Dataset contains {len(data)} records across {len(data.columns)} variables")
        
        # Missing data insights
        missing_pct = (data.isnull().sum() / len(data) * 100).round(2)
        high_missing = missing_pct[missing_pct > 10]
        if not high_missing.empty:
            insights.append(f"High missing data in columns: {list(high_missing.index)}")
        
        # Numeric insights
        numeric_cols = data.select_dtypes(include='number').columns
        if len(numeric_cols) > 0:
            insights.append(f"Numeric analysis available for {len(numeric_cols)} columns")
        
        return insights
    
    def _generate_recommendations(self, data: pd.DataFrame) -> List[str]:
        """Generate actionable recommendations"""
        recommendations = []
        
        # Data quality recommendations
        missing_pct = (data.isnull().sum() / len(data) * 100)
        if missing_pct.max() > 5:
            recommendations.append("Consider data cleaning for missing values")
        
        # Analysis recommendations
        if len(data.select_dtypes(include='number').columns) > 1:
            recommendations.append("Correlation analysis recommended for numeric variables")
        
        return recommendations
    
    def create_visualizations(self, data: pd.DataFrame, output_dir: str) -> List[str]:
        """Create data visualizations"""
        viz_paths = []
        
        try:
            # Ensure output directory exists
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            
            numeric_cols = data.select_dtypes(include='number').columns
            
            if len(numeric_cols) > 0:
                # Distribution plots
                for col in numeric_cols[:3]:  # Limit to first 3 columns
                    plt.figure(figsize=(10, 6))
                    sns.histplot(data[col].dropna())
                    plt.title(f'Distribution of {col}')
                    path = os.path.join(output_dir, f'dist_{col}.png')
                    plt.savefig(path, dpi=300, bbox_inches='tight')
                    plt.close()
                    viz_paths.append(path)
                
                # Correlation heatmap if multiple numeric columns
                if len(numeric_cols) > 1:
                    plt.figure(figsize=(12, 8))
                    corr_matrix = data[numeric_cols].corr()
                    sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', center=0)
                    plt.title('Correlation Heatmap')
                    path = os.path.join(output_dir, 'correlation_heatmap.png')
                    plt.savefig(path, dpi=300, bbox_inches='tight')
                    plt.close()
                    viz_paths.append(path)
            
            return viz_paths
        
        except Exception as e:
            logger.error(f"Visualization creation failed: {str(e)}")
            return []

class WordDocumentBuilder:
    """Handles Word document creation and formatting"""
    
    def __init__(self):
        self.document = None
        self.styles_applied = False
    
    def create_document(self, content: Dict[str, Any], request: DocumentRequest, 
                      analysis_results: Optional[Dict] = None, 
                      viz_paths: Optional[List[str]] = None) -> str:
        """Create a complete Word document"""
        
        self.document = Document()
        self._setup_styles()
        
        # Add content based on document type
        if request.doc_type == 'report':
            self._build_report(content, analysis_results, viz_paths)
        elif request.doc_type == 'article':
            self._build_article(content)
        elif request.doc_type == 'memo':
            self._build_memo(content)
        else:
            self._build_generic_document(content)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{request.doc_type}_{request.topic.replace(' ', '_')}_{timestamp}.docx"
        self.document.save(filename)
        
        return filename
    
    def _setup_styles(self):
        """Setup document styles"""
        if self.styles_applied:
            return
        
        # Create custom styles
        styles = self.document.styles
        
        # Heading styles
        if 'Custom Heading 1' not in [s.name for s in styles]:
            heading1 = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1.font.name = 'Calibri'
            heading1.font.size = Pt(18)
            heading1.font.bold = True
        
        self.styles_applied = True
    
    def _build_report(self, content: Dict[str, Any], analysis_results: Dict, viz_paths: List[str]):
        """Build a comprehensive report"""
        # Title page
        title = self.document.add_heading(content.get('title', 'Analysis Report'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive summary
        self.document.add_heading('Executive Summary', level=1)
        self.document.add_paragraph(content.get('executive_summary', 'Executive summary content'))
        
        # Data analysis section
        if analysis_results:
            self.document.add_heading('Data Analysis', level=1)
            
            # Summary statistics
            if 'summary' in analysis_results:
                summary = analysis_results['summary']
                p = self.document.add_paragraph()
                p.add_run(f"Dataset Overview: {summary.get('rows', 0)} records, {summary.get('columns', 0)} variables")
            
            # Insights
            if 'insights' in analysis_results:
                self.document.add_heading('Key Insights', level=2)
                for insight in analysis_results['insights']:
                    self.document.add_paragraph(insight, style='List Bullet')
            
            # Recommendations
            if 'recommendations' in analysis_results:
                self.document.add_heading('Recommendations', level=2)
                for rec in analysis_results['recommendations']:
                    self.document.add_paragraph(rec, style='List Bullet')
        
        # Add visualizations
        if viz_paths:
            self.document.add_heading('Data Visualizations', level=1)
            for viz_path in viz_paths:
                if os.path.exists(viz_path):
                    self.document.add_picture(viz_path, width=Inches(6))
                    self.document.add_paragraph()  # Add space
    
    def _build_article(self, content: Dict[str, Any]):
        """Build an article document"""
        # Title
        self.document.add_heading(content.get('title', 'Article Title'), 0)
        
        # Introduction
        self.document.add_heading('Introduction', level=1)
        self.document.add_paragraph(content.get('introduction', 'Introduction content'))
        
        # Main content
        self.document.add_heading('Main Content', level=1)
        self.document.add_paragraph(content.get('main_content', 'Main content'))
        
        # Conclusion
        self.document.add_heading('Conclusion', level=1)
        self.document.add_paragraph(content.get('conclusion', 'Conclusion content'))
    
    def _build_memo(self, content: Dict[str, Any]):
        """Build a memo document"""
        # Header
        header = content.get('header', 'MEMORANDUM')
        self.document.add_paragraph(header).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Memo details
        self.document.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        self.document.add_paragraph("To: [Recipients]")
        self.document.add_paragraph("From: [Sender]")
        self.document.add_paragraph("Subject: [Subject]")
        self.document.add_paragraph()  # Blank line
        
        # Content sections
        for section in ['purpose', 'summary', 'details', 'action_items']:
            if section in content:
                self.document.add_heading(section.replace('_', ' ').title(), level=1)
                self.document.add_paragraph(content[section])
    
    def _build_generic_document(self, content: Dict[str, Any]):
        """Build a generic document"""
        for section, text in content.items():
            self.document.add_heading(section.replace('_', ' ').title(), level=1)
            self.document.add_paragraph(text)

class WordAutomationAgent:
    """Main automation agent orchestrating all components"""
    
    def __init__(self, ai_api_key: str = None):
        self.content_generator = AIContentGenerator(ai_api_key)
        self.data_analyzer = DataAnalyzer()
        self.document_builder = WordDocumentBuilder()
        self.logger = logging.getLogger(__name__)
    
    def process_request(self, request: DocumentRequest) -> str:
        """Process a complete document generation request"""
        try:
            self.logger.info(f"Processing {request.doc_type} request: {request.topic}")
            
            # Generate AI content
            content = self.content_generator.generate_content(request)
            
            # Analyze data if provided
            analysis_results = None
            viz_paths = []
            
            if request.data_sources:
                for data_source in request.data_sources:
                    if os.path.exists(data_source):
                        analysis_results = self.data_analyzer.analyze_data(data_source)
                        
                        if request.include_charts:
                            data = self.data_analyzer._load_data(data_source)
                            viz_paths.extend(
                                self.data_analyzer.create_visualizations(data, 'output/charts')
                            )
            
            # Create document
            output_path = self.document_builder.create_document(
                content, request, analysis_results, viz_paths
            )
            
            self.logger.info(f"Document created successfully: {output_path}")
            return output_path
        
        except Exception as e:
            self.logger.error(f"Request processing failed: {str(e)}")
            raise
    
    def batch_process(self, requests: List[DocumentRequest]) -> List[str]:
        """Process multiple requests in batch"""
        results = []
        for request in requests:
            try:
                result = self.process_request(request)
                results.append(result)
            except Exception as e:
                self.logger.error(f"Batch processing failed for {request.topic}: {str(e)}")
                results.append(f"Error: {str(e)}")
        
        return results

# Example usage and configuration
def main():
    """Example usage of the Word Automation Agent"""
    
    # Initialize the agent
    agent = WordAutomationAgent(ai_api_key="your-openai-api-key-here")
    
    # Create a sample request
    sample_request = DocumentRequest(
        doc_type="report",
        topic="Sales Performance Analysis Q4 2024",
        length="long",
        tone="formal",
        audience="Executive Team",
        data_sources=["sample_data.csv"],  # Add your data file path
        include_charts=True,
        template="corporate_report"
    )
    
    # Process the request
    try:
        output_file = agent.process_request(sample_request)
        print(f"Document generated successfully: {output_file}")
    except Exception as e:
        print(f"Error generating document: {str(e)}")

# Additional utility functions
class ConfigurationManager:
    """Manage system configuration"""
    
    def __init__(self, config_file: str = "config.json"):
        self.config_file = config_file
        self.config = self.load_config()
    
    def load_config(self) -> Dict[str, Any]:
        """Load configuration from file"""
        if os.path.exists(self.config_file):
            with open(self.config_file, 'r') as f:
                return json.load(f)
        return self.get_default_config()
    
    def get_default_config(self) -> Dict[str, Any]:
        """Get default configuration"""
        return {
            "ai_model": "gpt-3.5-turbo",
            "output_directory": "output",
            "template_directory": "templates",
            "chart_style": "seaborn",
            "document_formats": ["docx", "pdf"],
            "max_file_size": "10MB",
            "supported_data_formats": [".csv", ".xlsx", ".json"],
        }
    
    def save_config(self):
        """Save current configuration"""
        with open(self.config_file, 'w') as f:
            json.dump(self.config, f, indent=2)

if __name__ == "__main__":
    main()
